"""
Bria Agentic Marketing Pipeline
================================
Brief → FIBO Generate → AI Review → FIBO Edit → Product Placement → Image Selection → Ads Template → Ad Heading Review

Run modes:
    python bria_marketing_agent.py generate --brief brief.txt         # Steps 1-5, outputs candidates
    python bria_marketing_agent.py finalize --candidates output/candidates.json --selected edited_0 edited_2  # Steps 7-8
    python bria_marketing_agent.py run --brief brief.txt              # Full pipeline, auto-select all
    python bria_marketing_agent.py --help                             # See all options

Requirements:
    pip install aiohttp anthropic python-dotenv python-docx
"""

import asyncio
import aiohttp
import base64
import json
import os
import sys
import argparse
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
from dotenv import load_dotenv

load_dotenv(override=True)

# ──────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────

BRIA_BASE_URL = "https://engine.prod.bria-api.com/v2"
BRIA_API_KEY = os.environ.get("BRIA_API_KEY", "")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
IMGBB_API_KEY = os.environ.get("IMGBB_API_KEY", "")

MAX_EDIT_RETRIES = 3
REVIEW_PASS_THRESHOLD = 8


# ──────────────────────────────────────────────
# Utilities — resolve local files to base64
# ──────────────────────────────────────────────

def is_url(val: str) -> bool:
    return val.strip().startswith("http://") or val.strip().startswith("https://")


def file_to_base64(path: str) -> str:
    """Read a local image file and return base64 string."""
    expanded = os.path.expanduser(path.strip())
    if not os.path.exists(expanded):
        raise FileNotFoundError(f"Image not found: {expanded}")
    with open(expanded, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def resolve_for_bria(path_or_url: str) -> str:
    """Return a URL string or base64 string suitable for Bria API image fields."""
    if is_url(path_or_url):
        return path_or_url.strip()
    return file_to_base64(path_or_url)


def validate_image_input(val: str) -> bool:
    """Check that the input is a valid URL or existing local file."""
    if is_url(val):
        return True
    return os.path.exists(os.path.expanduser(val.strip()))


# ──────────────────────────────────────────────
# Data Models
# ──────────────────────────────────────────────

@dataclass
class BriefSpec:
    campaign_name: str = ""
    brand_colors: list = field(default_factory=list)
    brand_fonts: list = field(default_factory=list)
    tone: str = ""
    target_audience: str = ""
    aspect_ratios: list = field(default_factory=lambda: ["1:1"])
    num_variants: int = 2
    visual_style: str = "photograph"
    scene_descriptions: list = field(default_factory=list)
    copy_text: list = field(default_factory=list)          # kept for backward compat
    ad_copy: list = field(default_factory=list)             # structured: [{"role":"heading","text":"..."}, ...]
    product_info: str = ""
    brand_id: Optional[str] = None
    template_ids: list = field(default_factory=list)


@dataclass
class ReviewResult:
    score: int = 0
    passed: bool = False
    feedback: str = ""
    edit_instructions: list = field(default_factory=list)
    dimensions: dict = field(default_factory=dict)


@dataclass
class PipelineResult:
    campaign_name: str = ""
    generated_images: list = field(default_factory=list)
    edited_images: list = field(default_factory=list)
    product_placed_images: list = field(default_factory=list)
    ad_creatives: list = field(default_factory=list)
    structured_prompts: list = field(default_factory=list)
    review_history: list = field(default_factory=list)
    ad_reviews: list = field(default_factory=list)


@dataclass
class UserInputs:
    brief_text: str = ""
    brief_file_path: Optional[str] = None                   # uploaded PDF/DOCX path
    brief_file_type: Optional[str] = None                   # ".pdf", ".docx", ".txt"
    product_images: list = field(default_factory=list)       # local paths or URLs
    reference_images: list = field(default_factory=list)     # local paths or URLs
    brand_id: Optional[str] = None
    template_ids: list = field(default_factory=list)
    output_dir: str = "output"
    bria_api_key: Optional[str] = None                      # override .env key from UI
    tailored_model_id: Optional[str] = None                 # fine-tuned model for generation
    tailored_model_influence: Optional[float] = None        # 0.0-1.0 weight for tailored model


# ──────────────────────────────────────────────
# Input Collection — From CLI arguments
# ──────────────────────────────────────────────

def collect_inputs_from_args(args) -> UserInputs:
    inputs = UserInputs()
    if getattr(args, 'brief_text', None):
        inputs.brief_text = args.brief_text
    elif getattr(args, 'brief', None):
        fp = os.path.expanduser(args.brief)
        if not os.path.exists(fp):
            print(f"❌ Not found: {fp}"); sys.exit(1)
        ext = os.path.splitext(fp)[1].lower()
        if ext in (".pdf", ".docx"):
            inputs.brief_file_path = fp
            inputs.brief_file_type = ext
        else:
            with open(fp) as f:
                inputs.brief_text = f.read()
    elif getattr(args, 'sample', False):
        inputs.brief_text = SAMPLE_BRIEF
    else:
        print("❌ No brief. Use --brief, --brief-text, or --sample."); sys.exit(1)
    if getattr(args, 'products', None):
        for p in args.products:
            if validate_image_input(p):
                inputs.product_images.append(p)
            else:
                print(f"⚠️ Product not found: {p}")
    if getattr(args, 'references', None):
        for r in args.references:
            if validate_image_input(r):
                inputs.reference_images.append(r)
            else:
                print(f"⚠️ Reference not found: {r}")
    inputs.brand_id = getattr(args, 'brand_id', None)
    inputs.template_ids = getattr(args, 'templates', None) or []
    inputs.output_dir = getattr(args, 'output', None) or "output"
    inputs.tailored_model_id = getattr(args, 'tailored_model', None)
    inputs.tailored_model_influence = getattr(args, 'tailored_influence', None)
    return inputs


# ──────────────────────────────────────────────
# Bria API Client — Async by default for generation, sync for edits
# ──────────────────────────────────────────────

class BriaClient:
    """
    Bria API client.
    Generation uses async mode (fire requests, poll for results).
    Editing and other operations use sync=true for simplicity.
    """

    def __init__(self, api_key: str, session: aiohttp.ClientSession):
        self.api_key = api_key
        self.session = session
        self.headers = {"api_token": api_key, "Content-Type": "application/json"}

    async def _post(self, endpoint: str, payload: dict, retries: int = 3, timeout: aiohttp.ClientTimeout | None = None) -> dict:
        """POST with sync=true. Retries on transient errors (401 SSL, 5xx)."""
        payload["sync"] = True
        url = f"{BRIA_BASE_URL}{endpoint}"

        for attempt in range(retries):
            try:
                async with self.session.post(url, headers=self.headers, json=payload, timeout=timeout) as resp:
                    data = await resp.json()
                    if resp.status == 200:
                        return data
                    # Transient errors — retry
                    if resp.status in (401, 500, 502, 503, 504) and attempt < retries - 1:
                        wait = 2 ** (attempt + 1)
                        print(f"   ⏳ Retry {attempt+1}/{retries} in {wait}s (HTTP {resp.status})...")
                        await asyncio.sleep(wait)
                        continue
                    raise Exception(f"Bria {resp.status}: {json.dumps(data)[:300]}")
            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                if attempt < retries - 1:
                    wait = 2 ** (attempt + 1)
                    print(f"   ⏳ Retry {attempt+1}/{retries} in {wait}s ({e})...")
                    await asyncio.sleep(wait)
                else:
                    raise

    async def _post_async(self, endpoint: str, payload: dict, retries: int = 3) -> dict:
        """POST without sync=true — returns immediately with request_id & status_url.
        Use _poll_status() to wait for the result."""
        url = f"{BRIA_BASE_URL}{endpoint}"

        for attempt in range(retries):
            try:
                async with self.session.post(url, headers=self.headers, json=payload) as resp:
                    data = await resp.json()
                    if resp.status in (200, 202):
                        return data
                    if resp.status in (401, 500, 502, 503, 504) and attempt < retries - 1:
                        wait = 2 ** (attempt + 1)
                        print(f"   ⏳ Retry {attempt+1}/{retries} in {wait}s (HTTP {resp.status})...")
                        await asyncio.sleep(wait)
                        continue
                    raise Exception(f"Bria {resp.status}: {json.dumps(data)[:300]}")
            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                if attempt < retries - 1:
                    wait = 2 ** (attempt + 1)
                    print(f"   ⏳ Retry {attempt+1}/{retries} in {wait}s ({e})...")
                    await asyncio.sleep(wait)
                else:
                    raise

    def _get_url(self, data: dict) -> str:
        """Extract image URL from various response shapes."""
        return (data.get("result", {}).get("image_url", "")
                or data.get("image_url", "")
                or data.get("result_url", ""))

    # ── Generation ──

    async def generate_image(self, prompt, aspect_ratio="1:1", num_results=1,
                             structured_prompt=None,
                             tailored_model_id=None,
                             tailored_model_influence=None):
        payload = {"prompt": prompt, "aspect_ratio": aspect_ratio, "num_results": num_results, "resolution": "4MP"}
        if structured_prompt:
            payload["structured_prompt"] = structured_prompt
        if tailored_model_id:
            payload["tailored_model_id"] = tailored_model_id
            if tailored_model_influence is not None:
                payload["model_influence"] = tailored_model_influence
        data = await self._post("/image/generate", payload)
        return self._get_url(data), data

    async def generate_inspired(self, prompt, reference_image, aspect_ratio="1:1"):
        """FIBO Inspire — reference_image can be local path or URL."""
        data = await self._post("/image/generate", {
            "prompt": prompt,
            "image": resolve_for_bria(reference_image),
            "aspect_ratio": aspect_ratio,
            "resolution": "4MP",
        })
        return self._get_url(data), data

    # ── Async Generation (fire-and-forget, poll later) ──

    async def generate_image_async(self, prompt, aspect_ratio="1:1", num_results=1,
                                   structured_prompt=None,
                                   tailored_model_id=None,
                                   tailored_model_influence=None) -> dict:
        """Fire generation request without waiting. Returns dict with status_url."""
        payload = {"prompt": prompt, "aspect_ratio": aspect_ratio,
                   "num_results": num_results, "resolution": "4MP"}
        if structured_prompt:
            payload["structured_prompt"] = structured_prompt
        if tailored_model_id:
            payload["tailored_model_id"] = tailored_model_id
            if tailored_model_influence is not None:
                payload["model_influence"] = tailored_model_influence
        return await self._post_async("/image/generate", payload)

    async def generate_inspired_async(self, prompt, reference_image, aspect_ratio="1:1") -> dict:
        """Fire inspired generation request without waiting. Returns dict with status_url."""
        return await self._post_async("/image/generate", {
            "prompt": prompt,
            "image": resolve_for_bria(reference_image),
            "aspect_ratio": aspect_ratio,
            "resolution": "4MP",
        })

    async def poll_batch(self, jobs: list[dict], max_wait: int = 300, interval: int = 5) -> list[dict]:
        """Poll multiple status_urls concurrently until all complete or fail.
        Each job dict must have 'status_url'. Returns completed result dicts
        in the same order, with '_error' key on failures."""

        async def _poll_one(job: dict) -> dict:
            status_url = job.get("status_url", "")
            if not status_url:
                return {**job, "_error": "No status_url in response"}
            try:
                result = await self._poll_status(status_url, max_wait=max_wait, interval=interval)
                return {**job, **result}
            except Exception as e:
                return {**job, "_error": str(e)}

        return await asyncio.gather(*[_poll_one(j) for j in jobs])

    # ── Editing ──

    async def edit_image(self, image_url, instruction):
        edit_timeout = aiohttp.ClientTimeout(total=30)
        data = await self._post("/image/edit", {
            "images": [image_url], "instruction": instruction
        }, timeout=edit_timeout)
        return self._get_url(data), data

    # ── Background ──

    async def remove_background(self, image):
        """image can be local path or URL."""
        data = await self._post("/image/edit/remove_background", {
            "image": resolve_for_bria(image)
        })
        return self._get_url(data), data

    # ── Product Placement ──
    # eCommerce product endpoints are at /v1/product/

    PRODUCT_BASE = "https://engine.prod.bria-api.com/v1/product"

    async def _poll_status(self, status_url: str, max_wait: int = 300, interval: int = 5) -> dict:
        """Poll a Bria status URL until terminal state (COMPLETED/ERROR/UNKNOWN).
        Works for both v1 and v2 status endpoints."""
        elapsed = 0
        while elapsed < max_wait:
            try:
                async with self.session.get(status_url, headers=self.headers) as resp:
                    data = await resp.json()
                    status = data.get("status", "").upper()

                    if status == "COMPLETED":
                        return data
                    elif status in ("ERROR", "UNKNOWN"):
                        err_msg = data.get("error", data.get("message", "Unknown error"))
                        raise Exception(f"Bria job {status}: {err_msg}")
                    # Still IN_PROGRESS — keep polling
                    print(f"      ⏳ {status} ({elapsed}s elapsed)...")
            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                print(f"      ⏳ Poll error ({e}), retrying...")

            await asyncio.sleep(interval)
            elapsed += interval

        raise Exception(f"Bria job timed out after {max_wait}s")

    async def _download_immediately(self, url: str) -> bytes:
        """Download image bytes immediately before signed URL expires.
        Tries with API headers first, then without."""
        # Try with API token headers (some Bria URLs need them)
        async with self.session.get(url, headers=self.headers) as r:
            if r.status == 200:
                return await r.read()
        # Fallback: try without headers (public URLs)
        async with self.session.get(url) as r:
            if r.status == 200:
                return await r.read()
            raise Exception(f"Download failed: HTTP {r.status} (tried with and without auth)")

    async def _post_product(self, endpoint: str, payload: dict, retries: int = 3) -> dict:
        """POST to product API. Tries sync=false with polling first;
        if the resulting URLs are not downloadable, falls back to sync=true.
        After getting the result, immediately downloads the image and adds
        '_image_bytes' to the response dict so callers can save it locally."""
        url = f"{self.PRODUCT_BASE}/{endpoint}"

        # Strategy 1: async (sync=false) — faster, but URLs may be signed/restricted
        async_payload = {**payload, "sync": False}
        for attempt in range(retries):
            try:
                async with self.session.post(url, headers=self.headers, json=async_payload) as resp:
                    raw = await resp.text()
                    if resp.status in (200, 202):
                        try:
                            data = json.loads(raw)
                        except json.JSONDecodeError:
                            raise Exception(f"Non-JSON response: {raw[:200]}")

                        # Poll if needed
                        if isinstance(data, dict):
                            status_url = data.get("status_url", "")
                            request_id = data.get("request_id", "")
                            if status_url:
                                print(f"      📡 Async job submitted (id: {request_id[:12]}...), polling...")
                                data = await self._poll_status(status_url)

                        # Try to download the image immediately
                        img_url = self._get_url_from_product(data)
                        if img_url:
                            try:
                                img_bytes = await self._download_immediately(img_url)
                                if not isinstance(data, dict):
                                    data = {"result": data}
                                data["_image_bytes"] = img_bytes
                                print(f"      ✓ Downloaded {len(img_bytes)} bytes")
                                return data
                            except Exception as de:
                                print(f"      ⚠️ Async URL not downloadable ({de}), trying sync mode...")
                                break  # Fall through to sync strategy
                        return data

                    if resp.status in (500, 502, 503, 504) and attempt < retries - 1:
                        wait = 3 * (attempt + 1)
                        print(f"   ⏳ Retry {attempt+1}/{retries} in {wait}s (HTTP {resp.status})...")
                        await asyncio.sleep(wait)
                        continue
                    raise Exception(f"Bria {resp.status}: {raw[:300]}")
            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                if attempt < retries - 1:
                    wait = 3 * (attempt + 1)
                    print(f"   ⏳ Retry {attempt+1}/{retries} in {wait}s ({e})...")
                    await asyncio.sleep(wait)
                else:
                    raise

        # Strategy 2: sync=true with long timeout — URLs are directly downloadable
        print(f"      🔄 Falling back to sync mode (longer timeout)...")
        sync_payload = {**payload, "sync": True}
        long_timeout = aiohttp.ClientTimeout(total=300)
        for attempt in range(retries):
            try:
                async with self.session.post(url, headers=self.headers, json=sync_payload,
                                              timeout=long_timeout) as resp:
                    raw = await resp.text()
                    if resp.status == 200:
                        try:
                            data = json.loads(raw)
                        except json.JSONDecodeError:
                            raise Exception(f"Non-JSON response: {raw[:200]}")

                        # Download immediately
                        img_url = self._get_url_from_product(data)
                        if img_url:
                            try:
                                img_bytes = await self._download_immediately(img_url)
                                if not isinstance(data, dict):
                                    data = {"result": data}
                                data["_image_bytes"] = img_bytes
                                print(f"      ✓ Downloaded {len(img_bytes)} bytes (sync)")
                            except Exception as de:
                                print(f"      ⚠️ Sync download failed: {de}")
                        return data

                    if resp.status in (500, 502, 503, 504) and attempt < retries - 1:
                        wait = 5 * (attempt + 1)
                        print(f"   ⏳ Sync retry {attempt+1}/{retries} in {wait}s (HTTP {resp.status})...")
                        await asyncio.sleep(wait)
                        continue
                    raise Exception(f"Bria {resp.status}: {raw[:300]}")
            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                if attempt < retries - 1:
                    wait = 5 * (attempt + 1)
                    print(f"   ⏳ Sync retry {attempt+1}/{retries} in {wait}s ({e})...")
                    await asyncio.sleep(wait)
                else:
                    raise

    async def lifestyle_shot_by_text(self, product_image_url, scene_description,
                                      num_results=2):
        """Place product in scene by text description."""
        data = await self._post_product("lifestyle_shot_by_text", {
            "image_url": product_image_url,
            "mode": "high_control",
            "scene_description": scene_description,
            "placement_type": "original",
            "num_results": num_results,
            "original_quality": True,
            "optimize_description": True,
        })
        return self._get_url_from_product(data), data

    async def lifestyle_shot_by_image(self, product_image_url, reference_image_url,
                                       num_results=2):
        """Place product in scene using reference image. ref_image_urls must be a list."""
        # Ensure ref is a list
        if isinstance(reference_image_url, str):
            ref_list = [reference_image_url]
        else:
            ref_list = list(reference_image_url)

        data = await self._post_product("lifestyle_shot_by_image", {
            "image_url": product_image_url,
            "ref_image_urls": ref_list,
            "placement_type": "original",
            "num_results": num_results,
            "original_quality": True,
        })
        return self._get_url_from_product(data), data

    def _get_url_from_product(self, data) -> str:
        """Extract image URL from product placement response.
        Handles various shapes:
          - {"result": [["url1","url2"]]}  (nested list of URLs)
          - {"result": [{"image_url":"..."}]}  (list of dicts)
          - {"result": {"image_url":"..."}}  (dict)
          - [{"image_url":"..."}]  (direct list of dicts)
          - [["url1"]]  (direct nested list)
        """
        def _extract_from_list(lst):
            """Recursively find first URL string in a possibly nested list."""
            for item in lst:
                if isinstance(item, str) and item.startswith("http"):
                    return item
                if isinstance(item, list):
                    found = _extract_from_list(item)
                    if found:
                        return found
                if isinstance(item, dict):
                    url = item.get("image_url", "") or item.get("url", "") or item.get("result_url", "")
                    if url:
                        return url
            return ""

        if isinstance(data, dict):
            result = data.get("result", data)
            if isinstance(result, list):
                return _extract_from_list(result)
            if isinstance(result, dict):
                return (result.get("image_url", "") or result.get("url", "")
                        or result.get("result_url", ""))
            if isinstance(result, str) and result.startswith("http"):
                return result
        if isinstance(data, list):
            return _extract_from_list(data)
        return ""

    # ── Transforms ──

    async def enhance(self, image_url):
        data = await self._post("/image/edit/enhance", {"image": image_url})
        return self._get_url(data), data

    # ── Ads ──
    # Ads API is at /v1/ads (not /v2)

    ADS_BASE = "https://engine.prod.bria-api.com/v1/ads"

    async def list_templates(self):
        async with self.session.get(f"{self.ADS_BASE}/templates",
                                     headers=self.headers) as r:
            return await r.json()

    async def get_template(self, tid):
        async with self.session.get(f"{self.ADS_BASE}/templates/{tid}",
                                     headers=self.headers) as r:
            data = await r.json()
            return data

    def extract_text_slots(self, template_data) -> list:
        """Extract valid text content_type values from a template response."""
        slots = []
        print(f"   [DEBUG] Template keys: {list(template_data.keys()) if isinstance(template_data, dict) else type(template_data).__name__}")
        print(f"   [DEBUG] Template preview: {json.dumps(template_data)[:600]}")

        if not isinstance(template_data, dict):
            return slots

        # Search all list values in the response for text layers
        def scan_dict(d, depth=0):
            if depth > 5:
                return
            for key, val in d.items():
                if isinstance(val, list):
                    for item in val:
                        if isinstance(item, dict):
                            lt = item.get("layer_type", "")
                            ct = item.get("content_type", "")
                            if lt == "text" and ct and ct not in slots:
                                slots.append(ct)
                            t = item.get("type", "")
                            n = item.get("name", "")
                            if t == "text" and n and n not in slots:
                                slots.append(n)
                            scan_dict(item, depth + 1)
                elif isinstance(val, dict):
                    scan_dict(val, depth + 1)

        scan_dict(template_data)
        return slots

    async def _poll_ad_url(self, ad_url: str, max_wait: int = 150, interval: int = 5) -> bytes:
        """Poll an ad image URL until it returns image data (HTTP 200).
        Ad images may take up to ~2 minutes to become available."""
        elapsed = 0
        while elapsed < max_wait:
            try:
                async with self.session.get(ad_url) as r:
                    if r.status == 200:
                        ct = r.headers.get("Content-Type", "")
                        if "image" in ct or "octet-stream" in ct:
                            return await r.read()
                    print(f"      ⏳ Ad image not ready (HTTP {r.status}, {elapsed}s elapsed)...")
            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                print(f"      ⏳ Ad poll error ({e}), retrying...")
            await asyncio.sleep(interval)
            elapsed += interval
        raise Exception(f"Ad image not available after {max_wait}s")

    @staticmethod
    def _extract_ad_url(data):
        """Extract first URL from ads response (kept for backward compat)."""
        urls = BriaClient._extract_ad_urls(data)
        return urls[0] if urls else ""

    @staticmethod
    def _extract_ad_urls(data):
        """Extract ALL URLs from ads response: {"result": [{"url": "..."}, ...]} etc."""
        urls = []

        def _extract_url_from_item(item):
            if isinstance(item, dict):
                return item.get("url", "") or item.get("image_url", "") or item.get("result_url", "")
            if isinstance(item, str) and item.startswith("http"):
                return item
            return ""

        if isinstance(data, dict):
            result = data.get("result", data)
            if isinstance(result, list):
                for item in result:
                    url = _extract_url_from_item(item)
                    if url:
                        urls.append(url)
            elif isinstance(result, dict):
                url = _extract_url_from_item(result)
                if url:
                    urls.append(url)
            if not urls:
                url = _extract_url_from_item(data)
                if url:
                    urls.append(url)
        elif isinstance(data, list):
            for item in data:
                url = _extract_url_from_item(item)
                if url:
                    urls.append(url)
        return urls

    async def create_ad(self, tid, brand_id, image_url, elements=None,
                        retries=3):
        """
        POST /v1/ads/generate with sync=true and expand_image operation.
        After getting the URLs, polls each until the image is populated (up to ~2min).
        Returns (list_of_urls, data_with_all_image_bytes).
        Each template may return a different number of ad variants.
        """
        payload = {
            "template_id": str(tid),
            "brand_id": str(brand_id),
            "sync": True,
            "smart_image": {
                "input_image_url": image_url,
                "scene": {
                    "operation": "expand_image"
                }
            },
            "elements": elements or []
        }

        url = f"{self.ADS_BASE}/generate"
        print(f"   [DEBUG] Ads payload: {json.dumps(payload)[:400]}")

        for attempt in range(retries):
            try:
                async with self.session.post(url, headers=self.headers, json=payload,
                                              timeout=aiohttp.ClientTimeout(total=180)) as resp:
                    content_type = resp.headers.get("Content-Type", "")
                    raw_text = await resp.text()

                    if resp.status in (200, 202):
                        try:
                            data = json.loads(raw_text)
                        except json.JSONDecodeError:
                            raise Exception(f"Ads returned non-JSON: {raw_text[:200]}")

                        print(f"   [DEBUG] Ads response: {json.dumps(data)[:300]}")

                        ad_urls = self._extract_ad_urls(data)
                        if not isinstance(data, dict):
                            data = {"result": data}
                        print(f"      📡 Found {len(ad_urls)} ad(s), waiting for images...")

                        # Download all ad images
                        all_bytes = []
                        for i, ad_url in enumerate(ad_urls):
                            try:
                                ad_bytes = await self._poll_ad_url(ad_url)
                                all_bytes.append(ad_bytes)
                                print(f"      ✓ Ad {i+1}/{len(ad_urls)} ready ({len(ad_bytes)} bytes)")
                            except Exception as e:
                                print(f"      ⚠️ Ad {i+1}/{len(ad_urls)} failed: {e}")
                                all_bytes.append(None)

                        data["_all_image_bytes"] = all_bytes
                        # Keep backward compat: _image_bytes = first image
                        if all_bytes and all_bytes[0]:
                            data["_image_bytes"] = all_bytes[0]
                        return ad_urls, data

                    if resp.status in (500, 502, 503, 504) and attempt < retries - 1:
                        wait = 5 * (attempt + 1)
                        print(f"   ⏳ Ads retry {attempt+1}/{retries} in {wait}s (HTTP {resp.status})...")
                        await asyncio.sleep(wait)
                        continue

                    print(f"   [DEBUG] Ads HTTP {resp.status}")
                    print(f"   [DEBUG] Content-Type: {content_type}")
                    print(f"   [DEBUG] Body: {raw_text[:500]}")
                    raise Exception(f"Ads {resp.status}: {raw_text[:300]}")

            except (aiohttp.ClientError, asyncio.TimeoutError) as e:
                if attempt < retries - 1:
                    wait = 5 * (attempt + 1)
                    print(f"   ⏳ Ads retry {attempt+1}/{retries} in {wait}s ({e})...")
                    await asyncio.sleep(wait)
                else:
                    raise


# ──────────────────────────────────────────────
# Image Downloader
# ──────────────────────────────────────────────

async def download(session, url, path):
    try:
        async with session.get(url) as r:
            if r.status == 200:
                path.parent.mkdir(parents=True, exist_ok=True)
                path.write_bytes(await r.read())
                return path
    except Exception as e:
        print(f"   ⚠️ Download failed: {e}")
    return None


async def upload_to_imgbb(session, image_url: str) -> str:
    """Download image from URL and re-upload to imgbb for a public link."""
    if not IMGBB_API_KEY:
        raise Exception("IMGBB_API_KEY not set in .env")

    # Download image bytes
    async with session.get(image_url) as r:
        if r.status != 200:
            raise Exception(f"Failed to download image: HTTP {r.status}")
        image_bytes = await r.read()

    # Upload to imgbb as base64
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    upload_url = "https://api.imgbb.com/1/upload"
    form = aiohttp.FormData()
    form.add_field("key", IMGBB_API_KEY)
    form.add_field("image", b64)

    async with session.post(upload_url, data=form) as r:
        data = await r.json()
        if r.status != 200 or not data.get("success"):
            raise Exception(f"imgbb upload failed: {json.dumps(data)[:200]}")
        public_url = data["data"]["url"]
        return public_url


async def upload_local_to_imgbb(session, local_path: str) -> str:
    """Upload a local image file to imgbb for a public link."""
    if not IMGBB_API_KEY:
        raise Exception("IMGBB_API_KEY not set in .env")

    with open(local_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    upload_url = "https://api.imgbb.com/1/upload"
    form = aiohttp.FormData()
    form.add_field("key", IMGBB_API_KEY)
    form.add_field("image", b64)

    async with session.post(upload_url, data=form) as r:
        data = await r.json()
        if r.status != 200 or not data.get("success"):
            raise Exception(f"imgbb upload failed: {json.dumps(data)[:200]}")
        return data["data"]["url"]


# ──────────────────────────────────────────────
# LLM Agent (Claude)
# ──────────────────────────────────────────────

class AgentLLM:

    def __init__(self, api_key, session):
        self.api_key = api_key
        self.session = session

    async def _call(self, system, user_content, max_tokens=2000, extra_headers=None):
        payload = {
            "model": "claude-sonnet-4-20250514",
            "max_tokens": max_tokens,
            "system": system,
            "messages": [{"role": "user", "content": user_content}]
        }
        headers = {
            "x-api-key": self.api_key, "anthropic-version": "2023-06-01",
            "Content-Type": "application/json"
        }
        if extra_headers:
            headers.update(extra_headers)
        async with self.session.post("https://api.anthropic.com/v1/messages",
                                     headers=headers, json=payload) as resp:
            data = await resp.json()
            if "error" in data:
                raise Exception(f"Claude: {data['error']}")
            return data["content"][0]["text"]

    def _json(self, text):
        text = text.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1].rsplit("```", 1)[0]
        return json.loads(text)

    async def parse_brief(self, brief_text="", brief_file_path=None, brief_file_type=None):
        system = """Parse marketing brief. Return ONLY valid JSON:
        {"campaign_name":"","brand_colors":["#hex"],"brand_fonts":[""],
        "tone":"","target_audience":"","aspect_ratios":["1:1"],
        "num_variants":2,"visual_style":"photograph",
        "scene_descriptions":["scene 1"],
        "copy_text":["headline1","subheading1"],
        "ad_copy":[
            {"role":"heading","text":"Main Headline"},
            {"role":"sub_heading","text":"Subheading text"},
            {"role":"cta","text":"Shop Now"}
        ],
        "product_info":""}

        IMPORTANT - ad_copy roles:
        - "heading": primary headline text for ads
        - "sub_heading": secondary/supporting text
        - "cta": call-to-action button text
        Classify each piece of copy text into the right role.
        copy_text is kept as a flat list of all text for backward compatibility.
        Default aspect_ratios to ["1:1"] if unspecified. Always extract ≥1 scene."""

        if brief_file_path and brief_file_type == ".pdf":
            # Anthropic native PDF support — send raw bytes as document block
            with open(brief_file_path, "rb") as f:
                pdf_b64 = base64.b64encode(f.read()).decode("utf-8")
            user_content = [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_b64,
                    },
                },
                {"type": "text", "text": "Parse this marketing brief document and extract all campaign details."},
            ]
            return BriefSpec(**self._json(await self._call(
                system, user_content,
                extra_headers={"anthropic-beta": "pdfs-2024-09-25"})))

        elif brief_file_path and brief_file_type in (".docx", ".doc"):
            # Extract text from DOCX using python-docx
            import docx
            doc = docx.Document(brief_file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n".join(paragraphs)
            return BriefSpec(**self._json(await self._call(system,
                [{"type": "text", "text": text}])))

        elif brief_file_path and brief_file_type == ".txt":
            with open(brief_file_path, "r") as f:
                text = f.read()
            return BriefSpec(**self._json(await self._call(system,
                [{"type": "text", "text": text}])))

        else:
            # Plain text from textarea (existing behavior)
            return BriefSpec(**self._json(await self._call(system,
                [{"type": "text", "text": brief_text}])))

    async def enrich_prompt(self, scene, spec):
        """Enrich a scene description with brand guidelines into a detailed prompt
        that Bria's VLM Bridge will convert to VGL internally."""
        system = """You are an expert image prompt writer. Take the scene description
        and brand context, and write a single detailed image generation prompt (1-3 sentences).
        Include: the scene, lighting, mood, color palette, composition, and style.
        Return ONLY the prompt text, no JSON, no markdown, no quotes."""
        result = await self._call(system, [{"type": "text", "text":
            f"Scene: {scene}\nBrand colors: {spec.brand_colors}\n"
            f"Tone: {spec.tone}\nAudience: {spec.target_audience}\n"
            f"Style: {spec.visual_style}"}], max_tokens=300)
        return result.strip()

    async def review(self, image_url, spec, scene, enriched_prompt="", product_info=""):
        system = """You are a strict creative director reviewing AI-generated marketing images.
Evaluate the image against the brief with rigorous, professional standards.

Score each dimension 1-10, then compute a weighted overall score.

SCORING DIMENSIONS:
1. scene_accuracy (25%): Does the image match the requested scene?
   Are all specified objects, setting, and environment elements present and correct?
   Score 1-3 if key elements are missing or wrong. 4-6 if partially correct.
   7-8 if mostly accurate. 9-10 ONLY if scene is exactly as described.

2. object_physics (20%): Are objects physically plausible?
   Check: objects resting on surfaces correctly, proper shadows, no floating objects,
   correct relative sizes, hands holding things naturally, liquids behaving correctly.
   Score 1-3 for obvious physics violations (floating, wrong scale, impossible poses).
   4-6 for subtle issues. 7-10 for physically convincing.

3. composition (15%): Is the image well-composed for marketing use?
   Subject placement, visual hierarchy, negative space for text overlay,
   not cluttered, not empty, focal point is clear.

4. brand_alignment (20%): Does the image match the brand's visual identity?
   Color palette matches brand colors, tone/mood matches brand tone,
   visual style matches requested style, appropriate for target audience.

5. technical_quality (10%): Is the image technically well-rendered?
   No artifacts, blur, distortion, anatomical errors, or border problems.

6. marketing_impact (10%): Would this be effective in a marketing campaign?
   Attention-grabbing, emotionally engaging, tells a clear story.

OVERALL = scene_accuracy*0.25 + object_physics*0.20 + composition*0.15 +
          brand_alignment*0.20 + technical_quality*0.10 + marketing_impact*0.10
Round to nearest integer.

PASS: overall >= 8 AND every dimension >= 5.
If ANY dimension < 5, the image FAILS regardless of overall score.

EDIT INSTRUCTIONS (if not passed):
Provide 1-3 specific, actionable instructions for the Bria FIBO image edit API.
Each must describe a concrete visual change — NOT abstract feedback.

GOOD: "Move the coffee cup so it rests flat on the table surface instead of floating"
GOOD: "Add warm golden sunlight streaming from the left side of the frame"
GOOD: "Change the background color from blue to warm cream to match brand palette"
BAD: "Fix the composition" (too vague for an edit API)
BAD: "Make it more on-brand" (not actionable)

Return ONLY valid JSON:
{
  "dimensions": {
    "scene_accuracy": {"score": 7, "reason": "brief explanation"},
    "object_physics": {"score": 4, "reason": "coffee cup appears to float beside the hand rather than being held"},
    "composition": {"score": 6, "reason": "brief explanation"},
    "brand_alignment": {"score": 7, "reason": "brief explanation"},
    "technical_quality": {"score": 8, "reason": "brief explanation"},
    "marketing_impact": {"score": 6, "reason": "brief explanation"}
  },
  "score": 6,
  "passed": false,
  "feedback": "one sentence overall assessment",
  "edit_instructions": ["specific FIBO edit instruction 1", "specific instruction 2"]
}"""

        # Build rich context for the review
        ctx = [f"Scene description: {scene}"]
        if enriched_prompt:
            ctx.append(f"Generation prompt used: {enriched_prompt}")
        ctx.append(f"Brand colors: {spec.brand_colors}")
        ctx.append(f"Tone/mood: {spec.tone}")
        ctx.append(f"Target audience: {spec.target_audience}")
        ctx.append(f"Visual style: {spec.visual_style}")
        if product_info or getattr(spec, 'product_info', ''):
            ctx.append(f"Product info: {product_info or spec.product_info}")

        result = self._json(await self._call(system, [
            {"type": "image", "source": {"type": "url", "url": image_url}},
            {"type": "text", "text": "\n".join(ctx)}
        ], max_tokens=1500))

        return ReviewResult(
            score=result.get("score", 0),
            passed=result.get("passed", False),
            feedback=result.get("feedback", ""),
            edit_instructions=result.get("edit_instructions", []),
            dimensions=result.get("dimensions", {}),
        )

    async def review_ad_headings(self, image_url, expected_texts):
        """Review an ad image to check if all headings are properly placed
        within the image borders and visually correct. Returns structured feedback."""
        system = """You are a QA reviewer for ad creatives. Analyze the ad image and check:
1. Are ALL text headings fully visible and inside the image borders (not cut off)?
2. Is each heading properly positioned (not overlapping other elements awkwardly)?
3. Is the text readable (not too small, not obscured by background)?

Return ONLY valid JSON:
{
  "passed": true/false,
  "issues": [
    {
      "text": "the heading text that has an issue",
      "problem": "description of the problem (e.g. 'cut off at bottom edge', 'overlaps product image')",
      "fix_instruction": "specific FIBO edit instruction to fix this (e.g. 'Move the heading text upward so it is fully visible within the image borders')"
    }
  ],
  "summary": "brief overall assessment"
}

If all headings are correctly placed, return passed=true with empty issues array.
Be strict — any text that is even partially cut off or outside borders should fail."""
        text_list = ", ".join(f'"{t}"' for t in expected_texts) if expected_texts else "unknown"
        return self._json(await self._call(system, [
            {"type": "image", "source": {"type": "url", "url": image_url}},
            {"type": "text", "text": f"Expected heading texts in this ad: {text_list}\n"
             "Check if all headings are fully visible, inside the image borders, "
             "and properly positioned. Report any issues."}], max_tokens=1500))


# ──────────────────────────────────────────────
# Pipeline
# ──────────────────────────────────────────────

class MarketingPipeline:

    def __init__(self, bria_key, anthropic_key, on_event=None):
        self.bria_key = bria_key
        self.anthropic_key = anthropic_key
        self.on_event = on_event  # Optional async callback: (event_type, data) -> None
        self.tailored_model_id = None    # Set by runner if user provides a fine-tuned model
        self.tailored_model_influence = None
        self.wait_for_selection = None   # Async callable set by runner for image selection pause

    async def _emit(self, event_type, data):
        """Emit event to callback if registered."""
        if self.on_event:
            try:
                await self.on_event(event_type, data)
            except Exception:
                pass  # Don't let event emission break the pipeline

    async def _run_generate_review_edit(self, spec, inputs, res, bria, agent, out, session):
        """Steps 2-4: Generate images (async batch), review them, and edit based on feedback."""
        # ═══ 2: Generate (async — fire all, then poll) ═══
        print("\n🎨 STEP 2: Generating (async mode — fire all, then poll)...")
        total_images = len(spec.scene_descriptions) * len(spec.aspect_ratios)
        await self._emit("step_start", {"step": 2, "step_id": "generate", "message": f"Generating {total_images} images ({len(spec.scene_descriptions)} scenes × {len(spec.aspect_ratios)} ratios)..."})
        has_refs = len(inputs.reference_images) > 0
        use_tailored = bool(self.tailored_model_id)

        if use_tailored:
            print(f"   🔧 Using tailored model: {self.tailored_model_id} (influence: {self.tailored_model_influence})")
            await self._emit("log", {"step": 2, "message": f"Using tailored model: {self.tailored_model_id} (influence: {self.tailored_model_influence})", "type": "info"})
            if has_refs:
                print(f"   ⚠️ Reference images are not used with tailored models — using text-to-image generation")
                await self._emit("log", {"step": 2, "message": "Note: Reference images are not used with tailored models — using text-to-image generation instead of Inspire", "type": "warning"})

        # ── Phase 1: Enrich all prompts ──
        enriched_prompts = []
        for i, scene in enumerate(spec.scene_descriptions):
            enriched_prompt = await agent.enrich_prompt(scene, spec)
            enriched_prompts.append(enriched_prompt)
            print(f"\n   Scene {i+1}: {scene[:55]}...")
            print(f"   Enriched: {enriched_prompt[:80]}...")
            await self._emit("log", {
                "step": 2,
                "message": f"Prompt for scene {i+1}: {enriched_prompt}",
                "type": "info",
                "subtype": "enriched_prompt",
            })

        # ── Phase 2: Fire all generation requests (no waiting) ──
        pending_jobs = []  # list of {scene_idx, scene, ratio, enriched_prompt, job_response}
        gen_count = 0
        for i, scene in enumerate(spec.scene_descriptions):
            enriched_prompt = enriched_prompts[i]
            for ratio in spec.aspect_ratios:
                gen_count += 1
                await self._emit("log", {"step": 2, "message": f"Firing request {gen_count}/{total_images}: scene {i+1} \"{scene[:50]}\" ({ratio})", "type": "info"})
                try:
                    if use_tailored:
                        print(f"   🚀 [{gen_count}/{total_images}] {ratio} via Tailored Model (async)...")
                        job = await bria.generate_image_async(
                            enriched_prompt, ratio, 1,
                            tailored_model_id=self.tailored_model_id,
                            tailored_model_influence=self.tailored_model_influence,
                        )
                    elif has_refs:
                        ref = inputs.reference_images[i % len(inputs.reference_images)]
                        print(f"   🚀 [{gen_count}/{total_images}] {ratio} via Inspire (async)...")
                        job = await bria.generate_inspired_async(enriched_prompt, ref, ratio)
                    else:
                        print(f"   🚀 [{gen_count}/{total_images}] {ratio} via FIBO (async)...")
                        job = await bria.generate_image_async(enriched_prompt, ratio, 1)

                    req_id = job.get("request_id", "???")[:12]
                    print(f"      📡 Submitted (id: {req_id}...)")
                    pending_jobs.append({
                        "scene_idx": i, "scene": scene, "ratio": ratio,
                        "enriched_prompt": enriched_prompt, **job,
                    })
                except Exception as e:
                    print(f"   ❌ Failed to submit: {e}")
                    await self._emit("log", {"step": 2, "message": f"Request failed for scene {i+1} ({ratio}): {e}", "type": "error"})

        if not pending_jobs:
            print("   ⚠️ No generation requests submitted")
            await self._emit("step_complete", {"step": 2, "step_id": "generate", "message": "No images generated"})
            return

        # ── Phase 3: Poll all status_urls concurrently ──
        print(f"\n   ⏳ Polling {len(pending_jobs)} jobs concurrently...")
        await self._emit("log", {"step": 2, "message": f"All {len(pending_jobs)} requests fired — polling for results...", "type": "info"})
        completed_jobs = await bria.poll_batch(pending_jobs)

        # ── Phase 4: Process results — download images ──
        for job in completed_jobs:
            i = job["scene_idx"]
            scene = job["scene"]
            ratio = job["ratio"]
            enriched_prompt = job["enriched_prompt"]

            if "_error" in job:
                print(f"   ❌ Scene {i+1} ({ratio}): {job['_error']}")
                await self._emit("log", {"step": 2, "message": f"Generation failed for scene {i+1} ({ratio}): {job['_error']}", "type": "error"})
                continue

            url = bria._get_url(job)
            returned_vgl = job.get("structured_prompt") or job.get("result", {}).get("structured_prompt")
            if returned_vgl:
                res.structured_prompts.append(returned_vgl)

            if url:
                fn = f"gen_s{i+1}_{ratio.replace(':','x')}.png"
                lp = await download(session, url, out / "1_generated" / fn)
                res.generated_images.append({
                    "scene": scene, "aspect_ratio": ratio,
                    "url": url, "local": str(lp),
                    "vgl": returned_vgl, "enriched_prompt": enriched_prompt})
                print(f"   ✓ 1_generated/{fn}")
                await self._emit("image_ready", {
                    "step": 2, "category": "generated", "url": url,
                    "scene": scene, "ratio": ratio, "local": str(lp),
                })
            else:
                print(f"   ⚠️ No URL in response for scene {i+1} ({ratio})")
                print(f"       Response keys: {list(job.keys())}")

        # ═══ 3-4: Review + Edit ═══
        print("\n🔍 STEP 3-4: Review & Edit...")
        await self._emit("step_complete", {"step": 2, "step_id": "generate", "message": f"Generated {len(res.generated_images)} images"})
        await self._emit("step_start", {"step": 3, "step_id": "review", "message": f"Reviewing {len(res.generated_images)} generated images..."})
        await self._emit("step_start", {"step": 4, "step_id": "edit", "message": "Editing images based on review feedback..."})
        for idx, img in enumerate(res.generated_images):
            url = img["url"]
            scene = img["scene"]
            ratio = img["aspect_ratio"]
            enriched_prompt = img.get("enriched_prompt", "")
            print(f"\n   [{idx+1}] {scene[:40]}... ({ratio})")
            await self._emit("log", {"step": 3, "message": f"Reviewing image {idx+1}/{len(res.generated_images)}: \"{scene[:50]}\"", "type": "info"})

            for att in range(MAX_EDIT_RETRIES + 1):
                try:
                    rv = await agent.review(
                        url, spec, scene,
                        enriched_prompt=enriched_prompt,
                        product_info=getattr(spec, 'product_info', ''),
                    )
                    res.review_history.append({
                        "scene": scene, "attempt": att,
                        "score": rv.score, "feedback": rv.feedback,
                        "dimensions": rv.dimensions,
                        "edit_instructions": rv.edit_instructions})
                    tag = "✅" if rv.passed else "⚠️"
                    print(f"   [{att+1}] {rv.score}/10 {tag} {rv.feedback[:70]}")
                    if rv.dimensions:
                        for dim_name, dim_data in rv.dimensions.items():
                            print(f"       {dim_name}: {dim_data.get('score', '?')}/10 — {dim_data.get('reason', '')[:50]}")
                    await self._emit("review_result", {
                        "step": 3, "scene": scene, "scene_index": idx,
                        "attempt": att, "score": rv.score, "passed": rv.passed,
                        "feedback": rv.feedback,
                        "dimensions": rv.dimensions,
                    })

                    if rv.passed or rv.score >= REVIEW_PASS_THRESHOLD:
                        fn = f"final_s{idx+1}_{ratio.replace(':','x')}.png"
                        lp = await download(session, url, out / "2_final" / fn)
                        res.edited_images.append({
                            **img, "final_url": url,
                            "local": str(lp), "score": rv.score})
                        await self._emit("image_ready", {
                            "step": 4, "category": "final", "url": url,
                            "scene": scene, "ratio": ratio, "local": str(lp),
                            "score": rv.score,
                        })
                        break

                    if att < MAX_EDIT_RETRIES:
                        for ei, instr in enumerate(rv.edit_instructions[:3]):
                            print(f"       ✏️ Edit {ei+1}: {instr[:55]}...")
                            await self._emit("log", {"step": 4, "message": f"Edit instruction {ei+1}/{len(rv.edit_instructions[:3])}: {instr}", "type": "info", "subtype": "edit_instruction"})
                            try:
                                url, _ = await bria.edit_image(url, instr)
                            except Exception as e:
                                print(f"       ⚠️ {e}")
                                await self._emit("log", {"step": 4, "message": f"Edit failed: {e}", "type": "warning"})
                    else:
                        print(f"   🚨 Flagged for human review")
                        res.edited_images.append({
                            **img, "final_url": url,
                            "score": rv.score, "flagged": True})
                        await self._emit("image_ready", {
                            "step": 4, "category": "final", "url": url,
                            "scene": scene, "ratio": ratio,
                            "score": rv.score, "flagged": True,
                        })
                except Exception as e:
                    print(f"   ❌ {e}")
                    res.edited_images.append({
                        **img, "final_url": url, "score": 0, "error": str(e)})
                    break

        await self._emit("step_complete", {"step": 3, "step_id": "review", "message": "Review complete"})
        await self._emit("step_complete", {"step": 4, "step_id": "edit", "message": f"Finalized {len(res.edited_images)} images"})

    async def _run_product_placement(self, spec, inputs, res, bria, out, session):
        """Step 5: Product placement."""
        if inputs.product_images:
            print(f"\n📦 STEP 5: Product Placement ({len(inputs.product_images)} products)...")
            await self._emit("step_start", {"step": 5, "step_id": "product", "message": f"Placing {len(inputs.product_images)} products into scenes..."})
            for pi, prod in enumerate(inputs.product_images):
                print(f"\n   Product {pi+1}: {prod}")
                await self._emit("log", {"step": 5, "message": f"Processing product {pi+1}/{len(inputs.product_images)}: removing background...", "type": "info"})
                try:
                    trans_url, _ = await bria.remove_background(prod)
                    print(f"   ✓ Background removed")

                    if IMGBB_API_KEY and trans_url:
                        print(f"   → Re-uploading to imgbb...")
                        trans_url = await upload_to_imgbb(session, trans_url)
                        print(f"   ✓ Public: {trans_url[:60]}...")
                except Exception as e:
                    print(f"   ❌ RMBG: {e}"); continue

                for si, scene in enumerate(spec.scene_descriptions):
                    try:
                        if inputs.reference_images:
                            ref = inputs.reference_images[si % len(inputs.reference_images)]
                            if not is_url(ref) and IMGBB_API_KEY:
                                ref_b64 = file_to_base64(ref)
                                form = aiohttp.FormData()
                                form.add_field("key", IMGBB_API_KEY)
                                form.add_field("image", ref_b64)
                                async with session.post("https://api.imgbb.com/1/upload", data=form) as r:
                                    rd = await r.json()
                                    if rd.get("success"):
                                        ref = rd["data"]["url"]
                            print(f"   → Scene {si+1} (ref image)...")
                            pu, pd = await bria.lifestyle_shot_by_image(trans_url, ref)
                        else:
                            print(f"   → Scene {si+1} (text)...")
                            pu, pd = await bria.lifestyle_shot_by_text(trans_url, scene)
                        if pu:
                            fn = f"prod{pi+1}_s{si+1}.png"
                            lp = out / "3_products" / fn
                            lp.parent.mkdir(parents=True, exist_ok=True)

                            img_bytes = pd.get("_image_bytes") if isinstance(pd, dict) else None
                            if img_bytes:
                                lp.write_bytes(img_bytes)
                            else:
                                lp = await download(session, pu, lp) or lp

                            public_pu = ""
                            if IMGBB_API_KEY and lp.exists():
                                try:
                                    public_pu = await upload_local_to_imgbb(session, str(lp))
                                    print(f"   ✓ Persisted: {public_pu[:60]}...")
                                except Exception as ue:
                                    print(f"   ⚠️ imgbb upload: {ue}")
                            res.product_placed_images.append({
                                "product": prod, "scene": scene,
                                "url": public_pu or pu, "local": str(lp)})
                            print(f"   ✓ 3_products/{fn}")
                            await self._emit("image_ready", {
                                "step": 5, "category": "product",
                                "url": public_pu or pu, "scene": scene,
                                "local": str(lp),
                            })
                    except Exception as e:
                        print(f"   ❌ {e}")
            await self._emit("step_complete", {"step": 5, "step_id": "product", "message": f"Placed {len(res.product_placed_images)} product images"})
        else:
            print("\n📦 STEP 5: Skipped")
            await self._emit("step_skipped", {"step": 5, "step_id": "product", "message": "No product images provided"})

    async def run(self, inputs: UserInputs, mode: str = "run") -> PipelineResult:
        """
        Run the marketing pipeline.
        mode:
          - "run": Full pipeline, auto-select all passing images for ads (Steps 1-8)
          - "generate": Steps 1-5 only, save candidates.json and exit
        """
        res = PipelineResult()
        out = Path(inputs.output_dir)
        out.mkdir(parents=True, exist_ok=True)

        async with aiohttp.ClientSession() as session:
            bria = BriaClient(self.bria_key, session)
            agent = AgentLLM(self.anthropic_key, session)

            # ═══ 1: Parse Brief ═══
            print("\n📋 STEP 1: Parsing brief...")
            await self._emit("step_start", {"step": 1, "step_id": "brief", "message": "Parsing marketing brief..."})
            spec = await agent.parse_brief(
                brief_text=inputs.brief_text,
                brief_file_path=inputs.brief_file_path,
                brief_file_type=inputs.brief_file_type,
            )
            res.campaign_name = spec.campaign_name
            if inputs.brand_id: spec.brand_id = inputs.brand_id
            if inputs.template_ids: spec.template_ids = inputs.template_ids

            print(f"   Campaign:  {spec.campaign_name}")
            print(f"   Tone:      {spec.tone}")
            print(f"   Colors:    {spec.brand_colors}")
            for i, s in enumerate(spec.scene_descriptions):
                print(f"   Scene {i+1}:   {s[:70]}...")
            print(f"   Formats:   {spec.aspect_ratios}")
            print(f"   Copy:      {spec.copy_text}")
            if spec.ad_copy:
                for ac in spec.ad_copy:
                    print(f"   Ad Copy:   [{ac.get('role', '?')}] {ac.get('text', '')}")
            print(f"   Brand ID:  {spec.brand_id!r} (from input: {inputs.brand_id!r})")
            print(f"   Templates: {spec.template_ids!r} (from input: {inputs.template_ids!r})")

            await self._emit("step_complete", {"step": 1, "step_id": "brief", "message": f"Brief parsed: {spec.campaign_name}"})
            await self._emit("brief_parsed", {"spec": {
                "campaign_name": spec.campaign_name, "brand_colors": spec.brand_colors,
                "tone": spec.tone, "target_audience": spec.target_audience,
                "scene_descriptions": spec.scene_descriptions, "aspect_ratios": spec.aspect_ratios,
                "copy_text": spec.copy_text,
                "brand_id": spec.brand_id,
                "template_ids": spec.template_ids,
            }})

            # ═══ 2-4: Generate + Review + Edit ═══
            await self._run_generate_review_edit(spec, inputs, res, bria, agent, out, session)

            # ═══ 5: Product Placement ═══
            await self._run_product_placement(spec, inputs, res, bria, out, session)

            # ═══ 6: Image Selection ═══
            # Build candidate list
            all_candidates = []
            for idx, img in enumerate(res.edited_images):
                all_candidates.append({
                    "id": f"edited_{idx}",
                    "category": "edited",
                    "url": img.get("final_url") or img.get("url"),
                    "local": img.get("local"),
                    "scene": img.get("scene", ""),
                    "score": img.get("score"),
                    "ratio": img.get("aspect_ratio", ""),
                })
            for idx, img in enumerate(res.product_placed_images):
                all_candidates.append({
                    "id": f"product_{idx}",
                    "category": "product",
                    "url": img.get("url"),
                    "local": img.get("local"),
                    "scene": img.get("scene", ""),
                })

            # Always save candidates.json for agent / external use
            candidates_file = out / "candidates.json"
            spec_dict = {
                "campaign_name": spec.campaign_name,
                "brand_colors": spec.brand_colors,
                "tone": spec.tone,
                "target_audience": spec.target_audience,
                "scene_descriptions": spec.scene_descriptions,
                "aspect_ratios": spec.aspect_ratios,
                "copy_text": spec.copy_text,
                "ad_copy": spec.ad_copy,
                "brand_id": spec.brand_id,
                "template_ids": spec.template_ids,
            }
            with open(candidates_file, "w") as f:
                json.dump({
                    "campaign_name": spec.campaign_name,
                    "candidates": all_candidates,
                    "spec": spec_dict,
                }, f, indent=2)
            print(f"\n📋 Saved {len(all_candidates)} candidates → {candidates_file}")

            # In "generate" mode, stop here — agent will handle selection
            if mode == "generate":
                print("\n✅ Generate phase complete. Candidates saved.")
                print(f"   Use 'finalize' command to continue with selected images.")
                # Still save partial pipeline_results.json
                self._save_results(res, inputs, spec, out)
                return res

            # In "run" mode, auto-select all images
            print(f"\n✅ STEP 6: Auto-selecting all {len(all_candidates)} images")
            await self._emit("step_complete", {"step": 6, "step_id": "image_selection", "message": f"Auto-selected {len(all_candidates)} images"})

            # ═══ 7: Ads ═══
            print(f"\n🖼️ STEP 7: Ads — brand_id={spec.brand_id!r}, template_ids={spec.template_ids!r}")
            if spec.brand_id and spec.template_ids:
                print(f"   → Proceeding with {len(spec.template_ids)} template(s)...")
                await self._emit("step_start", {"step": 7, "step_id": "ads", "message": f"Generating ad creatives with {len(spec.template_ids)} templates..."})
                # Process both product-placed images AND generated/final images
                heroes = res.product_placed_images + res.edited_images
                if not heroes:
                    msg = "No hero images available for ad generation"
                    print(f"   ⚠️ {msg}")
                    await self._emit("log", {"step": 7, "message": msg, "type": "warning"})
                    await self._emit("step_complete", {"step": 7, "step_id": "ads", "message": msg})
                elif not IMGBB_API_KEY:
                    msg = "IMGBB_API_KEY not set — ads step requires it for image hosting"
                    print(f"   ⚠️ {msg}")
                    await self._emit("log", {"step": 7, "message": msg, "type": "error"})
                    await self._emit("step_complete", {"step": 7, "step_id": "ads", "message": msg})
                else:
                    # Build role-based copy lookup from structured ad_copy
                    ad_copy_by_role = {}
                    for item in (spec.ad_copy or []):
                        role = item.get("role", "").lower()
                        if role and item.get("text"):
                            ad_copy_by_role.setdefault(role, []).append(item["text"])
                    # Flat fallback from copy_text
                    copy_texts_flat = [t for t in spec.copy_text if t]

                    for ti, tid in enumerate(spec.template_ids):
                        try:
                            await self._emit("log", {"step": 7, "message": f"Processing template {ti+1}/{len(spec.template_ids)} (ID: {tid})...", "type": "info"})
                            # Fetch template to discover valid text slots
                            tmpl = await bria.get_template(tid)
                            text_slots = bria.extract_text_slots(tmpl)
                            print(f"   Template {tid}: {len(text_slots)} text slot(s)")
                            for slot in text_slots:
                                print(f"     • {slot}")

                            # Map ad_copy roles to template slots by matching slot names
                            # Slot names like "heading", "sub_heading", "cta" etc.
                            elements = []
                            role_counters = {}  # track how many of each role we've used
                            flat_idx = 0        # fallback index into copy_texts_flat
                            for slot in text_slots:
                                slot_lower = slot.lower()
                                # Try to match slot name to an ad_copy role
                                matched_text = None
                                for role, texts in ad_copy_by_role.items():
                                    if role in slot_lower or slot_lower in role:
                                        idx = role_counters.get(role, 0)
                                        if idx < len(texts):
                                            matched_text = texts[idx]
                                            role_counters[role] = idx + 1
                                            break
                                # Fallback: use flat copy_text in order
                                if matched_text is None and flat_idx < len(copy_texts_flat):
                                    matched_text = copy_texts_flat[flat_idx]
                                    flat_idx += 1
                                if matched_text:
                                    elements.append({
                                        "layer_type": "text",
                                        "content_type": slot,
                                        "content": matched_text
                                    })

                            for img in heroes:
                                hero_url = img.get("url") or img.get("final_url")
                                local_path = img.get("local", "")

                                # Prefer uploading from local file (signed URLs may expire)
                                print(f"   → Uploading to imgbb...")
                                if local_path and os.path.exists(local_path):
                                    public_url = await upload_local_to_imgbb(session, local_path)
                                else:
                                    public_url = await upload_to_imgbb(session, hero_url)
                                print(f"   ✓ Public URL: {public_url[:60]}...")

                                print(f"   → Generating ad (expand_image)...")
                                ad_urls, ad_data = await bria.create_ad(
                                    tid=tid,
                                    brand_id=spec.brand_id,
                                    image_url=public_url,
                                    elements=elements,
                                )

                                # Save ALL ad images returned for this template
                                all_bytes = ad_data.get("_all_image_bytes", []) if isinstance(ad_data, dict) else []
                                for ad_idx, ad_url in enumerate(ad_urls):
                                    fn = f"ad_{tid}_{len(res.ad_creatives)+1}.png"
                                    ad_path = out / "4_ads" / fn
                                    ad_path.parent.mkdir(parents=True, exist_ok=True)

                                    ad_bytes = all_bytes[ad_idx] if ad_idx < len(all_bytes) else None
                                    if ad_bytes:
                                        ad_path.write_bytes(ad_bytes)
                                    else:
                                        await download(session, ad_url, ad_path)

                                    # Store the actual text elements sent to this template
                                    ad_texts = [e["content"] for e in elements if e.get("content")]
                                    res.ad_creatives.append({
                                        "tid": tid, "url": ad_url, "local": str(ad_path),
                                        "ad_texts": ad_texts,
                                        "elements": elements
                                    })
                                    print(f"   ✓ 4_ads/{fn}")
                                    await self._emit("image_ready", {
                                        "category": "ad",
                                        "url": f"/api/images/{str(ad_path)}",
                                        "scene": f"Ad – Template {tid} #{len(res.ad_creatives)}",
                                        "ratio": None,
                                        "local": str(ad_path),
                                    })

                                if not ad_urls:
                                    print(f"   ⚠️ No ad URLs returned for template {tid}")
                        except Exception as e:
                            print(f"   ❌ Template {tid} error: {e}")
                            await self._emit("log", {"step": 7, "message": f"Template {tid} failed: {e}", "type": "error"})
                    await self._emit("step_complete", {"step": 7, "step_id": "ads", "message": f"Generated {len(res.ad_creatives)} ads"})
            else:
                reason = []
                if not spec.brand_id:
                    reason.append(f"brand_id is empty ({spec.brand_id!r})")
                if not spec.template_ids:
                    reason.append(f"template_ids is empty ({spec.template_ids!r})")
                msg = f"Skipped — {'; '.join(reason)}"
                print(f"   {msg}")
                await self._emit("step_skipped", {"step": 7, "step_id": "ads", "message": msg})

            # ═══ 8: Ad Heading Review & Fix (SKIPPED) ═══
            print("\n🔎 STEP 8: Skipped (ad heading review disabled)")
            await self._emit("step_skipped", {"step": 8, "step_id": "ad_review", "message": "Ad heading review skipped"})

        # ═══ Save JSON & Summary ═══
        self._save_results(res, inputs, spec, out)
        return res

    def _save_results(self, res, inputs, spec, out):
        """Save pipeline_results.json and print summary."""
        out = Path(out)
        rp = out / "pipeline_results.json"
        with open(rp, "w") as f:
            json.dump({
                "campaign": res.campaign_name,
                "generated": [{"scene": g["scene"], "ratio": g["aspect_ratio"],
                               "url": g["url"], "local": g.get("local")} for g in res.generated_images],
                "final": [{"scene": e["scene"], "url": e["final_url"],
                           "score": e.get("score"), "local": e.get("local")} for e in res.edited_images],
                "products": [{"scene": p["scene"], "url": p["url"],
                              "local": p.get("local")} for p in res.product_placed_images],
                "ads": res.ad_creatives,
                "ad_reviews": res.ad_reviews,
                "reviews": res.review_history,
                "vgl_prompts": [p for p in res.structured_prompts if p],
                "enriched_prompts": [g.get("enriched_prompt", "") for g in res.generated_images],
                "_debug": {
                    "input_brand_id": inputs.brand_id if inputs else None,
                    "input_template_ids": inputs.template_ids if inputs else [],
                    "spec_brand_id": spec.brand_id if spec else None,
                    "spec_template_ids": spec.template_ids if spec else [],
                    "hero_count": len(res.product_placed_images + res.edited_images),
                    "imgbb_key_set": bool(IMGBB_API_KEY),
                }
            }, f, indent=2)

        print("\n" + "=" * 60)
        print(f"✅ DONE: {res.campaign_name}")
        print("=" * 60)
        print(f"   Generated:  {len(res.generated_images)}")
        print(f"   Final:      {len(res.edited_images)}")
        print(f"   Products:   {len(res.product_placed_images)}")
        print(f"   Ads:        {len(res.ad_creatives)}")
        ad_passed = sum(1 for r in res.ad_reviews if r.get("passed"))
        if res.ad_reviews:
            print(f"   Ad Reviews: {ad_passed}/{len(set(r.get('ad_index') for r in res.ad_reviews))} passed")
        print(f"\n   📂 {out.resolve()}")
        print("\n📁 Files:")
        for f in sorted(out.rglob("*.png")):
            print(f"   {f.relative_to(out)}")

    async def run_finalize(self, candidates_path: str, selected_ids: list,
                           brand_id: str = None, template_ids: list = None,
                           output_dir: str = None) -> PipelineResult:
        """
        Run Steps 7-8 (Ad Generation + Review) using previously generated candidates.
        Loads candidates.json, filters to selected images, generates ads.
        """
        with open(candidates_path) as f:
            data = json.load(f)

        spec_data = data.get("spec", {})
        candidates = data.get("candidates", [])
        out = Path(output_dir) if output_dir else Path(candidates_path).parent

        # Override brand_id / template_ids from CLI if provided
        if brand_id:
            spec_data["brand_id"] = brand_id
        if template_ids:
            spec_data["template_ids"] = template_ids

        # Build a minimal BriefSpec
        spec = BriefSpec()
        for k, v in spec_data.items():
            if hasattr(spec, k):
                setattr(spec, k, v)

        # Filter candidates to selected IDs
        selected_set = set(selected_ids)
        selected_candidates = [c for c in candidates if c["id"] in selected_set]
        if not selected_candidates:
            print("❌ No matching candidates for the selected IDs")
            print(f"   Available: {[c['id'] for c in candidates]}")
            print(f"   Selected:  {selected_ids}")
            sys.exit(1)

        print(f"\n✅ Loaded {len(selected_candidates)} selected images for ad generation")
        for c in selected_candidates:
            print(f"   • {c['id']}: {c.get('scene', '')[:60]}...")

        # Build a PipelineResult with the selected images
        res = PipelineResult()
        res.campaign_name = data.get("campaign_name", "Campaign")
        for c in selected_candidates:
            if c["category"] == "edited":
                res.edited_images.append({
                    "scene": c.get("scene", ""),
                    "url": c.get("url", ""),
                    "final_url": c.get("url", ""),
                    "local": c.get("local", ""),
                    "score": c.get("score"),
                    "aspect_ratio": c.get("ratio", ""),
                })
            elif c["category"] == "product":
                res.product_placed_images.append({
                    "scene": c.get("scene", ""),
                    "url": c.get("url", ""),
                    "local": c.get("local", ""),
                })

        # Run Steps 7-8
        async with aiohttp.ClientSession() as session:
            bria = BriaClient(self.bria_key, session)

            # ═══ 7: Ads ═══
            print(f"\n🖼️ STEP 7: Ads — brand_id={spec.brand_id!r}, template_ids={spec.template_ids!r}")
            if spec.brand_id and spec.template_ids:
                print(f"   → Proceeding with {len(spec.template_ids)} template(s)...")
                heroes = res.product_placed_images + res.edited_images
                if not heroes:
                    print("   ⚠️ No hero images available")
                elif not IMGBB_API_KEY:
                    print("   ⚠️ IMGBB_API_KEY not set — required for ad generation")
                else:
                    ad_copy_by_role = {}
                    for item in (spec.ad_copy or []):
                        role = item.get("role", "").lower()
                        if role and item.get("text"):
                            ad_copy_by_role.setdefault(role, []).append(item["text"])
                    copy_texts_flat = [t for t in spec.copy_text if t]

                    for ti, tid in enumerate(spec.template_ids):
                        try:
                            tmpl = await bria.get_template(tid)
                            text_slots = bria.extract_text_slots(tmpl)
                            print(f"   Template {tid}: {len(text_slots)} text slot(s)")

                            elements = []
                            role_counters = {}
                            flat_idx = 0
                            for slot in text_slots:
                                slot_lower = slot.lower()
                                matched_text = None
                                for role, texts in ad_copy_by_role.items():
                                    if role in slot_lower or slot_lower in role:
                                        idx = role_counters.get(role, 0)
                                        if idx < len(texts):
                                            matched_text = texts[idx]
                                            role_counters[role] = idx + 1
                                            break
                                if matched_text is None and flat_idx < len(copy_texts_flat):
                                    matched_text = copy_texts_flat[flat_idx]
                                    flat_idx += 1
                                if matched_text:
                                    elements.append({
                                        "layer_type": "text",
                                        "content_type": slot,
                                        "content": matched_text
                                    })

                            for img in heroes:
                                hero_url = img.get("url") or img.get("final_url")
                                local_path = img.get("local", "")

                                print(f"   → Uploading to imgbb...")
                                if local_path and os.path.exists(local_path):
                                    public_url = await upload_local_to_imgbb(session, local_path)
                                else:
                                    public_url = await upload_to_imgbb(session, hero_url)
                                print(f"   ✓ Public URL: {public_url[:60]}...")

                                print(f"   → Generating ad (expand_image)...")
                                ad_urls, ad_data = await bria.create_ad(
                                    tid=tid, brand_id=spec.brand_id,
                                    image_url=public_url, elements=elements,
                                )

                                all_bytes = ad_data.get("_all_image_bytes", []) if isinstance(ad_data, dict) else []
                                for ad_idx, ad_url in enumerate(ad_urls):
                                    fn = f"ad_{tid}_{len(res.ad_creatives)+1}.png"
                                    ad_path = out / "4_ads" / fn
                                    ad_path.parent.mkdir(parents=True, exist_ok=True)

                                    ad_bytes = all_bytes[ad_idx] if ad_idx < len(all_bytes) else None
                                    if ad_bytes:
                                        ad_path.write_bytes(ad_bytes)
                                    else:
                                        await download(session, ad_url, ad_path)

                                    ad_texts = [e["content"] for e in elements if e.get("content")]
                                    res.ad_creatives.append({
                                        "tid": tid, "url": ad_url, "local": str(ad_path),
                                        "ad_texts": ad_texts, "elements": elements
                                    })
                                    print(f"   ✓ 4_ads/{fn}")
                        except Exception as e:
                            print(f"   ❌ Template {tid} error: {e}")
            else:
                print("   Skipped — no brand_id or template_ids")

            # ═══ 8: Ad Heading Review (SKIPPED) ═══
            print("\n🔎 STEP 8: Skipped (ad heading review disabled)")

        # Save results
        self._save_results(res, None, spec, out)
        return res


# ──────────────────────────────────────────────
# Sample Brief
# ──────────────────────────────────────────────

SAMPLE_BRIEF = """
Campaign: Summer Refresh 2026
Brand: FreshBrew Coffee
Brand Colors: #2D5016 (forest green), #F5E6D3 (cream), #8B4513 (coffee brown)
Fonts: Playfair Display, Source Sans Pro
Tone: Warm, inviting, premium but approachable
Target Audience: Urban professionals 25-40
Required Formats: Instagram square (1:1)
Scenes:
1. Steaming coffee on a sunlit cafe table with pastries, morning light
2. Cold brew bottle on a picnic blanket in a park, summer afternoon
Headlines: "Refresh Your Ritual" / "Crafted for Your Best Moments"
"""


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────

def _add_common_args(parser):
    """Add arguments shared by 'generate' and 'run' subcommands."""
    parser.add_argument("--brief", "-b", help="Brief file path (txt, pdf, docx)")
    parser.add_argument("--brief-text", help="Brief as inline text string")
    parser.add_argument("--sample", "-s", action="store_true", help="Use built-in sample brief")
    parser.add_argument("--products", "-p", nargs="+", help="Product images (paths/URLs)")
    parser.add_argument("--references", "-r", nargs="+", help="Reference images (paths/URLs)")
    parser.add_argument("--brand-id", help="Bria Brand ID")
    parser.add_argument("--templates", "-t", nargs="+", help="Bria Template IDs")
    parser.add_argument("--output", "-o", default="output", help="Output directory")
    parser.add_argument("--tailored-model", help="Fine-tuned Bria model ID")
    parser.add_argument("--tailored-influence", type=float, help="Tailored model influence (0.0-1.0)")


def build_parser():
    p = argparse.ArgumentParser(
        description="Bria Agentic Marketing Pipeline",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python bria_marketing_agent.py generate --brief brief.txt -o output
  python bria_marketing_agent.py generate --brief-text "Campaign: ..." -o output
  python bria_marketing_agent.py finalize --candidates output/candidates.json --selected edited_0 edited_2
  python bria_marketing_agent.py finalize --candidates output/candidates.json --selected edited_0 --brand-id 162 --templates 1274
  python bria_marketing_agent.py run --brief brief.txt --brand-id 162 --templates 1274
  python bria_marketing_agent.py run --sample
        """)
    sub = p.add_subparsers(dest="command", help="Pipeline command")

    # generate: Steps 1-5, output candidates.json
    gen_p = sub.add_parser("generate", help="Run Steps 1-5, output candidates for selection")
    _add_common_args(gen_p)

    # finalize: Steps 7-8 with selected images
    fin_p = sub.add_parser("finalize", help="Run Steps 7-8 with selected images")
    fin_p.add_argument("--candidates", required=True, help="Path to candidates.json")
    fin_p.add_argument("--selected", nargs="+", required=True, help="Selected candidate IDs (e.g. edited_0 edited_2 product_0)")
    fin_p.add_argument("--brand-id", help="Bria Brand ID (overrides brief)")
    fin_p.add_argument("--templates", "-t", nargs="+", help="Bria Template IDs (overrides brief)")
    fin_p.add_argument("--output", "-o", help="Output directory (default: same as candidates)")

    # run: Full pipeline, auto-select all
    run_p = sub.add_parser("run", help="Run full pipeline (auto-select all passing images)")
    _add_common_args(run_p)

    return p


async def main():
    args = build_parser().parse_args()

    if not args.command:
        build_parser().print_help()
        sys.exit(1)

    if not BRIA_API_KEY:
        print("❌ BRIA_API_KEY missing — add to .env"); sys.exit(1)
    if not ANTHROPIC_API_KEY:
        print("❌ ANTHROPIC_API_KEY missing — add to .env"); sys.exit(1)
    if not IMGBB_API_KEY:
        print("⚠️  IMGBB_API_KEY not set — ad generation will be skipped")

    pipeline = MarketingPipeline(BRIA_API_KEY, ANTHROPIC_API_KEY)

    if args.command == "finalize":
        # Steps 7-8 only
        output_dir = args.output or str(Path(args.candidates).parent)
        await pipeline.run_finalize(
            candidates_path=args.candidates,
            selected_ids=args.selected,
            brand_id=args.brand_id,
            template_ids=args.templates,
            output_dir=output_dir,
        )
    else:
        # "generate" or "run"
        inputs = collect_inputs_from_args(args)
        if inputs.tailored_model_id:
            pipeline.tailored_model_id = inputs.tailored_model_id
        if inputs.tailored_model_influence is not None:
            pipeline.tailored_model_influence = inputs.tailored_model_influence
        await pipeline.run(inputs, mode=args.command)


if __name__ == "__main__":
    asyncio.run(main())
